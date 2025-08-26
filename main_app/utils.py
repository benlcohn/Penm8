from docx import Document
from bs4 import BeautifulSoup
import re
import nltk

try:
    from nltk.corpus import cmudict
    nltk.data.find("corpora/cmudict")
    cmu_dict = cmudict.dict()
except LookupError:
    cmu_dict = None


def count_syllables_in_word(word: str) -> int:
    word = word.lower()
    if cmu_dict and word in cmu_dict:
        # cmudict can give multiple pronunciations; take min syllable count
        return min(
            [len([ph for ph in pron if ph[-1].isdigit()]) for pron in cmu_dict[word]]
        )
    else:
        # Fallback heuristic
        word = re.sub(r"[^a-z]", "", word)
        if not word:
            return 0
        vowels = "aeiouy"
        count = 0
        prev_char_was_vowel = False
        for char in word:
            if char in vowels:
                if not prev_char_was_vowel:
                    count += 1
                prev_char_was_vowel = True
            else:
                prev_char_was_vowel = False
        # silent "e"
        if word.endswith("e"):
            count = max(1, count - 1)
        return max(1, count)


def process_docx(file_path: str) -> tuple[str, int, int, int, int, int, int]:
    """
    Read DOCX with python-docx, convert to HTML preserving:
    - Paragraph alignment
    - Inline styles (bold, italic, underline)
    - Multiple spaces/tabs
    - Blank lines (as <br> or empty <p>)
    Returns:
      (html_content, word_count, char_count, sentence_count, line_count, paragraph_count, syllable_count)
    """
    doc = Document(file_path)

    html_parts = []
    align_map = {
        0: "left",
        1: "center",
        2: "right",
        3: "justify",
    }

    for para in doc.paragraphs:
        if not para.text.strip():
            # Preserve blank lines
            html_parts.append("<p><br></p>")
            continue

        # Paragraph alignment
        align = align_map.get(para.paragraph_format.alignment, "left")
        para_html = f'<p style="text-align:{align}; white-space: pre-wrap;">'

        # Inline runs
        for run in para.runs:
            # Preserve line breaks *inside runs*
            text = run.text.replace("\n", "<br>")

            # Preserve multiple spaces & tabs
            text = text.replace("  ", "&nbsp;&nbsp;")   # double spaces
            text = text.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")  # tabs

            if not text:
                continue

            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"

            para_html += text

        para_html += "</p>"
        html_parts.append(para_html)

    # Final HTML
    html_content = "".join(html_parts)

    # Plain text for analysis
    soup = BeautifulSoup(html_content, "html.parser")
    plain_text = soup.get_text(separator="\n")  # keep line breaks for counting

    # Word & char counts
    words = plain_text.split()
    word_count = len(words)
    clean_text = plain_text.replace("\n", "").replace("\t", "")
    char_count = len(clean_text)

    # Sentence count
    sentences = re.split(r"[.!?]+(?:\s|$)", plain_text.strip())
    sentence_count = len([s for s in sentences if s.strip()])

    # Line count (based on newlines in plain text)
    line_count = len([line for line in plain_text.splitlines() if line.strip()])

    # Paragraph count (based on <p> tags we generated)
    paragraph_count = len(soup.find_all("p"))

    # Syllable count (total doc)
    syllable_count = sum(count_syllables_in_word(w) for w in words)

    return (
        html_content,
        word_count,
        char_count,
        sentence_count,
        line_count,
        paragraph_count,
        syllable_count,
    )

def process_html(html: str) -> tuple[str, int, int, int, int, int, int]:
    """
    Analyze already-formatted HTML (from CKEditor).
    Returns:
      (html_content, word_count, char_count, sentence_count,
       line_count, paragraph_count, syllable_count)
    """
    # Ensure valid HTML
    soup = BeautifulSoup(html, "html.parser")

    # Cleaned HTML (preserve inline tags, normalize whitespace)
    html_content = str(soup)

    # Extract plain text
    plain_text = soup.get_text(separator="\n")

    # Word & char counts
    words = plain_text.split()
    word_count = len(words)
    clean_text = plain_text.replace("\n", "").replace("\t", "")
    char_count = len(clean_text)

    # Sentence count
    sentences = re.split(r"[.!?]+(?:\s|$)", plain_text.strip())
    sentence_count = len([s for s in sentences if s.strip()])

    # Line count (based on newlines in plain text)
    line_count = len([line for line in plain_text.splitlines() if line.strip()])

    # Paragraph count (<p> tags)
    paragraph_count = len(soup.find_all("p"))

    # Syllable count
    syllable_count = sum(count_syllables_in_word(w) for w in words)

    return (
        html_content,
        word_count,
        char_count,
        sentence_count,
        line_count,
        paragraph_count,
        syllable_count,
    )

def process_docx_perline(file_path: str) -> list[dict]:
    """
    Analyze the DOCX line by line, returning a list of dicts:
    [
      {"text": "<b>Hello</b> world", "words": 2, "syllables": 3},
      ...
    ]

    - Preserves inline formatting (<b>, <i>, <u>)
    - Preserves blank lines as &nbsp;
    - Preserves leading spaces/tabs
    - Counts words/syllables from raw text (ignores tags)
    """
    doc = Document(file_path)

    line_stats = []
    for para in doc.paragraphs:
        if not para.text.strip():
            line_stats.append({"text": "&nbsp;", "words": -1, "syllables": -1})
            continue

        # Build inner HTML for the paragraph (like in process_docx)
        html_runs = []
        for run in para.runs:
            text = run.text.replace("\n", "<br>")
            text = text.replace("  ", "&nbsp;&nbsp;")
            text = text.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")

            if not text:
                continue
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"

            html_runs.append(text)

        raw_html = "".join(html_runs).strip()
        raw_text = para.text.strip()

        if raw_text == "":
            line_stats.append({"text": "&nbsp;", "words": -1, "syllables": -1})
            continue

        words = re.findall(r"\b\w+\b", raw_text)
        syllables = sum(count_syllables_in_word(w) for w in words)

        line_stats.append(
            {
                "text": raw_html,     # preserves inline formatting
                "words": len(words),
                "syllables": syllables,
            }
        )

    return line_stats

def process_html_perline(html_content: str) -> list[dict]:
    """
    Analyze HTML input (from CKEditor) line by line.
    Preserves <b>, <i>, <u>, etc.
    Blank lines get a &nbsp; placeholder.
    Counts words/syllables from the raw *text* (ignores tags).
    """
    soup = BeautifulSoup(html_content, "html.parser")

    line_stats = []
    for p in soup.find_all("p"):
        # Get the inner HTML of the <p> instead of plain text
        raw_html = "".join(str(c) for c in p.contents).strip()
        raw_text = p.get_text().strip()

        if raw_text == "":
            # Preserve blank lines visually with &nbsp;
            line_stats.append({"text": "&nbsp;", "words": -1, "syllables": -1})
            continue

        # Count words and syllables from raw *text* (not HTML)
        words = re.findall(r"\b\w+\b", raw_text)
        syllables = sum(count_syllables_in_word(w) for w in words)

        line_stats.append(
            {
                "text": raw_html,      # preserves <b>, <i>, <u>, etc.
                "words": len(words),
                "syllables": syllables,
            }
        )

    return line_stats

